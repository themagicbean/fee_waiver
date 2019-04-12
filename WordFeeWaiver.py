#! python3

'''
Created on Mar 27, 2019

WordFeeWaiver -- makes a fee waiver motion

@author: darrenbean
'''

import docx
from django.db.models.functions import text
doc = docx.Document("Form Fee Waiver.docx")

from tkinter import *
from tkinter import messagebox

master = Tk()
Label(master, text="First Name").grid(row=0)
Label(master, text="Last Name").grid(row=1)
Label(master, text="Case No.").grid(row=2)
Label(master, text="Date").grid(row=3)
Label(master, text="Time").grid(row=4)
Label(master, text="Dept.").grid(row=5)

DFIRST = Entry(master)
DFIRST.grid(row=0, column=1)
DLAST = Entry(master)
DLAST.grid(row=1, column=1)
CASENO = Entry(master)
CASENO.grid(row=2, column=1)
HRGDATE = Entry(master)
HRGDATE.grid(row=3, column=1)
HRGTIME = Entry(master)
HRGTIME.grid(row=4, column=1)
HRGDEPT = Entry(master)
HRGDEPT.grid(row=5, column=1)

def makefeewaiver():
    
    global DFIRST, DLAST, CASENO, HRGDATE, HRGTIME, HRGDEPT

    info = {"DFIRST": DFIRST.get(),
        "DLAST": DLAST.get(),
        "CASENO": CASENO.get(),
        "HRGDATE": HRGDATE.get(),
        "HRGTIME": HRGTIME.get(),
        "HRGDEPT": HRGDEPT.get()}  
    
    for p in doc.paragraphs:
        for key in info.keys():
            if key in p.text:
                text = p.text.replace(key, info[key])
                style = p.style
                p.text = text
                p.style = style
         
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key in info.keys():
                        if key in p.text:
                            text = p.text.replace(key, info[key])
                            style = p.style
                            p.text = text
                            p.style = style
    
    doc.save(str(info["DFIRST"]) + " " + str(info["DLAST"]) + " Fee Waiver.docx")
    messagebox.showinfo("Complete!", "The Motion has Been Made and Saved")
    
b1 = Button(master, relief="raised", text="Make Fee Waiver", command=makefeewaiver)
b1.grid(row=7)
b2 = Button(master, relief="raised", text="Cancel and Quit", command=master.quit)
b2.grid(row=8)

mainloop()


     
   


